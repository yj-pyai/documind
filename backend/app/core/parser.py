"""Document parsing: extract text from various file formats."""

import os


async def parse_document(filepath: str, file_type: str) -> str:
    """Parse a document file and extract plain text.

    Supported formats: pdf, docx, md, txt
    """
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "md": _parse_text,
        "txt": _parse_text,
        "markdown": _parse_text,
    }

    parser = parsers.get(file_type.lower())
    if parser is None:
        raise ValueError(f"Unsupported file type: {file_type}")

    return parser(filepath)


def _parse_pdf(filepath: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF (fitz) is required for PDF parsing. Install: pip install PyMuPDF")

    doc = fitz.open(filepath)
    text_parts = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text("text")
        if page_text.strip():
            text_parts.append(page_text)

    doc.close()
    return "\n\n".join(text_parts)


def _parse_docx(filepath: str) -> str:
    """Extract text from Word documents."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install: pip install python-docx")

    doc = Document(filepath)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            # Check if it's a heading
            if para.style.name.startswith("Heading"):
                level = para.style.name.split()[-1]
                try:
                    hashes = "#" * int(level)
                    text_parts.append(f"{hashes} {para.text.strip()}")
                except ValueError:
                    text_parts.append(para.text.strip())
            else:
                text_parts.append(para.text.strip())

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n\n".join(text_parts)


def _parse_text(filepath: str) -> str:
    """Parse plain text or markdown files."""
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()
